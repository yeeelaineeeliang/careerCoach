import streamlit as st
import anthropic
import json
import os
from pathlib import Path

# Load .env if present (so ANTHROPIC_API_KEY works when set in .env)
try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).parent / ".env", override=True)
except ImportError:
    pass
import sqlite3
from datetime import date, datetime
from pathlib import Path
from typing import Optional

# ─── PAGE CONFIG ───
st.set_page_config(
    page_title="CareerOS",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── CUSTOM CSS ───
with open(Path(__file__).parent / "styles.css") as _f:
    st.markdown(f"<style>{_f.read()}</style>", unsafe_allow_html=True)

# ─── DORAEMON BUDDY WIDGET ───
# CSS + HTML via st.markdown (styles render fine; script tags are stripped by Streamlit/React)
st.markdown("""
<div id="doraemon-float">
  <div id="doraemon-bubble">✨ Hey Annie! I'm cheering you on~ 💙</div>
  <div id="doraemon-avatar" title="Drag me anywhere · Click for a surprise!">
    <svg width="88" height="108" viewBox="0 0 200 240" xmlns="http://www.w3.org/2000/svg">
      <ellipse cx="100" cy="228" rx="50" ry="9" fill="rgba(0,0,0,0.18)"/>
      <ellipse cx="100" cy="172" rx="68" ry="58" fill="#0099DD"/>
      <ellipse cx="100" cy="180" rx="46" ry="40" fill="white"/>
      <ellipse cx="100" cy="181" rx="32" ry="18" fill="#0d1020" stroke="rgba(0,153,221,0.3)" stroke-width="1.5"/>
      <path d="M68 181 Q100 167 132 181" stroke="rgba(0,153,221,0.3)" stroke-width="1.5" fill="none"/>
      <ellipse cx="72"  cy="222" rx="28" ry="14" fill="white"/>
      <ellipse cx="128" cy="222" rx="28" ry="14" fill="white"/>
      <ellipse cx="38"  cy="172" rx="16" ry="10" fill="#0099DD" transform="rotate(-20 38 172)"/>
      <circle  cx="24"  cy="182" r="14" fill="white"/>
      <ellipse cx="162" cy="172" rx="16" ry="10" fill="#0099DD" transform="rotate(20 162 172)"/>
      <circle  cx="176" cy="182" r="14" fill="white"/>
      <g class="dora-tail">
        <path d="M148 200 Q168 195 172 210 Q170 222 158 218" stroke="#0099DD" stroke-width="8" fill="none" stroke-linecap="round"/>
        <circle cx="160" cy="218" r="8" fill="#0099DD"/>
        <circle cx="160" cy="218" r="4" fill="white"/>
      </g>
      <circle cx="100" cy="96"  r="72" fill="#0099DD"/>
      <ellipse cx="100" cy="108" rx="54" ry="46" fill="white"/>
      <ellipse cx="60"  cy="116" rx="16" ry="10" fill="#ff6666" opacity="0.55"/>
      <ellipse cx="140" cy="116" rx="16" ry="10" fill="#ff6666" opacity="0.55"/>
      <circle cx="78"  cy="88" r="16" fill="white" class="dora-eye"/>
      <circle cx="82"  cy="90" r="9"  fill="#111"/>
      <circle cx="85"  cy="87" r="3.5" fill="white"/>
      <circle cx="122" cy="88" r="16" fill="white" class="dora-eye"/>
      <circle cx="118" cy="90" r="9"  fill="#111"/>
      <circle cx="121" cy="87" r="3.5" fill="white"/>
      <circle cx="100" cy="104" r="9" fill="#cc2222"/>
      <ellipse cx="97" cy="101" rx="3" ry="2" fill="white" opacity="0.85"/>
      <line x1="100" y1="113" x2="100" y2="128" stroke="#aaa" stroke-width="2"/>
      <path id="dora-mouth" d="M72 128 Q100 150 128 128" stroke="#aaa" stroke-width="2.5" fill="none" stroke-linecap="round"/>
      <line x1="30"  y1="108" x2="88"  y2="114" stroke="#666" stroke-width="1.5" stroke-linecap="round"/>
      <line x1="28"  y1="120" x2="88"  y2="120" stroke="#666" stroke-width="1.5" stroke-linecap="round"/>
      <line x1="30"  y1="132" x2="88"  y2="126" stroke="#666" stroke-width="1.5" stroke-linecap="round"/>
      <line x1="170" y1="108" x2="112" y2="114" stroke="#666" stroke-width="1.5" stroke-linecap="round"/>
      <line x1="172" y1="120" x2="112" y2="120" stroke="#666" stroke-width="1.5" stroke-linecap="round"/>
      <line x1="170" y1="132" x2="112" y2="126" stroke="#666" stroke-width="1.5" stroke-linecap="round"/>
      <rect x="50" y="140" width="100" height="16" rx="8" fill="#cc2222"/>
      <circle cx="100" cy="154" r="10" fill="#FFCC00"/>
      <circle cx="100" cy="157" r="3"  fill="#aa8800"/>
      <line x1="91" y1="154" x2="109" y2="154" stroke="#aa8800" stroke-width="1.5"/>
      <circle cx="96" cy="150" r="2.5" fill="white" opacity="0" class="dora-bell"/>
    </svg>
  </div>
</div>
""", unsafe_allow_html=True)

# JavaScript is injected via components.html() which runs in an iframe on the same origin,
# allowing window.parent access — this is the correct way to run JS in Streamlit.
import streamlit.components.v1 as _st_components
_st_components.html("""
<script>
(function() {
  var pd = window.parent.document;

  var MSGS = [
    "You're doing amazing~ 💙 Keep it up!",
    "I believe in you, Annie! 🌟",
    "One step at a time~ I'm with you! ✨",
    "Every application gets you closer! 🎯",
    "You've totally got this! 💪💙",
    "Stay curious — that's your superpower! 🔧",
    "Dream big, grind smart~ 🚀",
    "I'd give you a gadget, but you're already amazing! 🎒",
    "Proud of you for showing up today! 🌟",
    "Your offer is coming~ just keep going! 💙",
    "Rest when you need to — I'll still be here! 🌙",
    "Small wins count! Celebrate every step~ 🎉",
    "You showed up today. That already matters! 💫",
  ];

  var AUTO_MSGS = [
    "✨ Hey Annie! Cheering you on~ 💙",
    "💙 You're making progress — keep going!",
    "🌟 Today's effort = tomorrow's offer!",
    "🎯 Focus mode: activated! You've got this~",
    "💫 I'm right here with you, always!",
    "🔵 Every rejection is just a redirect~ 💙",
  ];

  var EMOJIS = ['💙','⭐','✨','💫','🌟','🎯','💪'];

  function init() {
    var floatEl = pd.getElementById('doraemon-float');
    var avatar  = pd.getElementById('doraemon-avatar');
    var bubble  = pd.getElementById('doraemon-bubble');
    var mouth   = pd.getElementById('dora-mouth');

    if (!floatEl || !avatar || !bubble) { setTimeout(init, 120); return; }
    if (avatar.dataset.doraReady) return;   // prevent double-init on re-render
    avatar.dataset.doraReady = '1';

    /* ── Drag state ── */
    var dragging = false, moved = false;
    var startMouseX, startMouseY, startLeft, startBottom;

    avatar.addEventListener('mousedown', function(e) {
      if (e.button !== 0) return;
      dragging = true; moved = false;
      startMouseX = e.clientX; startMouseY = e.clientY;
      var rect = floatEl.getBoundingClientRect();
      startLeft   = rect.left;
      startBottom = window.parent.innerHeight - rect.bottom;
      avatar.classList.add('dora-dragging');
      e.preventDefault();
    });

    pd.addEventListener('mousemove', function(e) {
      if (!dragging) return;
      var dx = e.clientX - startMouseX;
      var dy = e.clientY - startMouseY;
      if (Math.abs(dx) > 4 || Math.abs(dy) > 4) moved = true;
      var newLeft   = Math.max(0, Math.min(window.parent.innerWidth  - 96,  startLeft   + dx));
      var newBottom = Math.max(0, Math.min(window.parent.innerHeight - 120, startBottom - dy));
      floatEl.style.right  = 'auto';
      floatEl.style.left   = newLeft   + 'px';
      floatEl.style.bottom = newBottom + 'px';
    });

    pd.addEventListener('mouseup', function(e) {
      if (!dragging) return;
      dragging = false;
      avatar.classList.remove('dora-dragging');
      if (!moved) tap(e);   // treat no-move mousedown+up as click
    });

    /* ── Touch drag ── */
    avatar.addEventListener('touchstart', function(e) {
      var t = e.touches[0];
      dragging = true; moved = false;
      startMouseX = t.clientX; startMouseY = t.clientY;
      var rect = floatEl.getBoundingClientRect();
      startLeft   = rect.left;
      startBottom = window.parent.innerHeight - rect.bottom;
    }, {passive: true});

    pd.addEventListener('touchmove', function(e) {
      if (!dragging) return;
      var t = e.touches[0];
      var dx = t.clientX - startMouseX;
      var dy = t.clientY - startMouseY;
      moved = true;
      var newLeft   = Math.max(0, Math.min(window.parent.innerWidth  - 96,  startLeft   + dx));
      var newBottom = Math.max(0, Math.min(window.parent.innerHeight - 120, startBottom - dy));
      floatEl.style.right  = 'auto';
      floatEl.style.left   = newLeft   + 'px';
      floatEl.style.bottom = newBottom + 'px';
    }, {passive: true});

    pd.addEventListener('touchend', function(e) {
      if (!dragging) return;
      dragging = false;
      if (!moved) {
        var t = e.changedTouches[0];
        tap({clientX: t.clientX, clientY: t.clientY});
      }
    });

    /* ── Tap / click ── */
    var lastMsgIdx = -1, hideTimer = null;

    function tap(e) {
      // Pick a new random message
      var idx;
      do { idx = Math.floor(Math.random() * MSGS.length); } while (idx === lastMsgIdx);
      lastMsgIdx = idx;

      // Show bubble
      bubble.classList.remove('dora-hidden');
      bubble.innerHTML = MSGS[idx];
      clearTimeout(hideTimer);
      hideTimer = setTimeout(function(){ bubble.classList.add('dora-hidden'); }, 5500);

      // Bounce animation on avatar
      avatar.classList.remove('dora-bounce');
      void avatar.offsetWidth;            // reflow to restart animation
      avatar.classList.add('dora-bounce');
      setTimeout(function(){ avatar.classList.remove('dora-bounce'); }, 400);

      // Widen smile briefly
      if (mouth) {
        mouth.setAttribute('d', 'M66 128 Q100 158 134 128');
        setTimeout(function(){ mouth.setAttribute('d', 'M72 128 Q100 150 128 128'); }, 600);
      }

      // Burst of particles
      var rect = floatEl.getBoundingClientRect();
      var cx = rect.left + rect.width  / 2;
      var cy = rect.top  + rect.height / 2;
      for (var i = 0; i < 6; i++) {
        (function(i){
          setTimeout(function(){
            var el = pd.createElement('div');
            el.className = 'dora-particle';
            el.textContent = EMOJIS[Math.floor(Math.random() * EMOJIS.length)];
            el.style.left = (cx + (Math.random()-0.5)*70) + 'px';
            el.style.top  = (cy + (Math.random()-0.5)*50) + 'px';
            pd.body.appendChild(el);
            setTimeout(function(){ el.remove(); }, 1200);
          }, i * 60);
        })(i);
      }
    }

    /* ── Auto-rotate idle messages every 14s ── */
    var autoIdx = 0;
    setInterval(function(){
      autoIdx = (autoIdx + 1) % AUTO_MSGS.length;
      bubble.classList.remove('dora-hidden');
      bubble.innerHTML = AUTO_MSGS[autoIdx];
      clearTimeout(hideTimer);
      hideTimer = setTimeout(function(){ bubble.classList.add('dora-hidden'); }, 6000);
    }, 14000);
  }

  init();
})();
</script>
""", height=0)

# ─── SQLITE PERSISTENCE ───
DB_PATH = Path(__file__).parent / "careeros.db"

def _db():
    con = sqlite3.connect(str(DB_PATH))
    con.execute("""
        CREATE TABLE IF NOT EXISTS kv (
            key   TEXT PRIMARY KEY,
            value TEXT NOT NULL
        )
    """)
    con.commit()
    return con

def db_save(key: str, value) -> None:
    con = _db()
    con.execute(
        "INSERT OR REPLACE INTO kv (key, value) VALUES (?, ?)",
        (key, json.dumps(value, default=str))
    )
    con.commit()
    con.close()

def db_load(key: str, default=None):
    try:
        con = _db()
        row = con.execute("SELECT value FROM kv WHERE key=?", (key,)).fetchone()
        con.close()
        return json.loads(row[0]) if row else default
    except Exception:
        return default

def db_load_all() -> dict:
    """Load all persisted keys into a single dict on startup."""
    try:
        con = _db()
        rows = con.execute("SELECT key, value FROM kv").fetchall()
        con.close()
        return {k: json.loads(v) for k, v in rows}
    except Exception:
        return {}

# ─── PROJECTS DATABASE ───
def _projects_db():
    con = sqlite3.connect(str(DB_PATH))
    con.execute("""
        CREATE TABLE IF NOT EXISTS projects (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            title         TEXT    NOT NULL,
            source_type   TEXT    NOT NULL,
            raw_content   TEXT    NOT NULL,
            technologies  TEXT    DEFAULT '[]',
            metrics       TEXT    DEFAULT '[]',
            contributions TEXT    DEFAULT '',
            challenges    TEXT    DEFAULT '',
            summary       TEXT    DEFAULT '',
            created_at    REAL    NOT NULL
        )
    """)
    con.commit()
    return con

def db_save_project(title: str, source_type: str, raw_content: str, extracted: dict) -> int:
    con = _projects_db()
    cur = con.execute(
        """INSERT INTO projects
           (title, source_type, raw_content, technologies, metrics, contributions, challenges, summary, created_at)
           VALUES (?,?,?,?,?,?,?,?,?)""",
        (
            title, source_type, raw_content[:60000],
            json.dumps(extracted.get("technologies", []), ensure_ascii=False),
            json.dumps(extracted.get("metrics", []), ensure_ascii=False),
            extracted.get("contributions", ""),
            extracted.get("challenges", ""),
            extracted.get("summary", ""),
            datetime.now().timestamp(),
        )
    )
    pid = cur.lastrowid
    con.commit()
    con.close()
    return pid

def db_get_projects() -> list:
    try:
        con = _projects_db()
        rows = con.execute(
            "SELECT id, title, source_type, technologies, metrics, contributions, challenges, summary, created_at "
            "FROM projects ORDER BY created_at DESC"
        ).fetchall()
        con.close()
        return [
            {
                "id": r[0], "title": r[1], "source_type": r[2],
                "technologies": json.loads(r[3] or "[]"),
                "metrics":      json.loads(r[4] or "[]"),
                "contributions": r[5] or "",
                "challenges":    r[6] or "",
                "summary":       r[7] or "",
                "created_at":    r[8],
            }
            for r in rows
        ]
    except Exception:
        return []

def db_delete_project(project_id: int) -> None:
    con = _projects_db()
    con.execute("DELETE FROM projects WHERE id=?", (project_id,))
    con.commit()
    con.close()

def read_pdf_bytes(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pypdf."""
    import io as _io
    try:
        import pypdf
        reader = pypdf.PdfReader(_io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                pages.append(text.strip())
        return "\n\n".join(pages)
    except Exception as e:
        return f"[PDF extraction error: {e}]"

def read_docx_bytes(file_bytes: bytes) -> str:
    """Extract text from a .docx file.
    Primary: python-docx. Fallback: stdlib zipfile + XML (no extra deps required).
    """
    import io as _io

    # ── Primary: python-docx ──
    try:
        import docx as _docx
        doc = _docx.Document(_io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = "  |  ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        if parts:
            return "\n\n".join(parts)
    except Exception:
        pass

    # ── Fallback: pure stdlib ZIP + XML parse ──
    try:
        import zipfile as _zf
        import xml.etree.ElementTree as _ET

        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        with _zf.ZipFile(_io.BytesIO(file_bytes)) as z:
            with z.open("word/document.xml") as f:
                tree = _ET.parse(f)

        root = tree.getroot()
        paragraphs = []
        for para in root.iter(f"{{{W}}}p"):
            runs = [node.text for node in para.iter(f"{{{W}}}t") if node.text]
            line = "".join(runs).strip()
            if line:
                paragraphs.append(line)

        if paragraphs:
            return "\n\n".join(paragraphs)
        return "[DOCX appears to have no readable text content]"
    except Exception as e:
        return f"[DOCX extraction error: {e}]"

def extract_project_with_claude(raw_text: str, title_hint: str = "") -> dict:
    """Call Claude Haiku to extract structured project data. Only uses facts present in the text."""
    import re as _re
    client = get_client()
    if not client or client == "invalid_format":
        return {
            "title": title_hint or "Untitled Project",
            "technologies": [], "metrics": [],
            "contributions": raw_text[:400],
            "challenges": "", "summary": raw_text[:200],
        }
    prompt = f"""You are extracting structured facts from a project report or description.
Extract ONLY what is explicitly stated — never invent or infer metrics or details not present in the text.

PROJECT TITLE HINT: {title_hint or "(infer from content)"}

PROJECT CONTENT:
{raw_text[:8000]}

Return a JSON object with EXACTLY these keys:
{{
  "title": "project title (inferred or confirmed)",
  "technologies": ["each specific tool / language / framework / library / platform explicitly mentioned"],
  "metrics": ["copy exact numbers/percentages/scale from the text only — empty array if none"],
  "contributions": "2-4 sentences on what this person specifically built, owned, or led (first-person perspective, past tense)",
  "challenges": "1-3 sentences on key technical problems solved or design decisions made",
  "summary": "2-sentence summary optimized for matching against job descriptions"
}}

STRICT RULES:
- metrics: only include figures EXPLICITLY in the text. If none exist, return [].
- Never add facts not in the source text.
- Return ONLY valid JSON. No markdown fences, no other text."""
    try:
        response = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=900,
            messages=[{"role": "user", "content": prompt}]
        )
        text = _extract_text_blocks(response.content)
        match = _re.search(r'\{.*\}', text, _re.DOTALL)
        if match:
            data = json.loads(match.group())
            if "title" in data:
                return data
    except Exception:
        pass
    return {
        "title": title_hint or "Untitled Project",
        "technologies": [], "metrics": [],
        "contributions": raw_text[:400],
        "challenges": "", "summary": raw_text[:200],
    }

def build_project_context_for_resume() -> str:
    """Format all saved projects into a context block for the Resume Reviewer prompt."""
    projects = db_get_projects()
    if not projects:
        return ""
    lines = [
        "━━━ PROJECT LIBRARY — Annie's Verified Work (ground truth — use these facts, never hallucinate) ━━━",
        "INSTRUCTION: When rewriting or evaluating resume bullets, pull relevant details from the projects below.",
        "Auto-match: identify which project(s) best align with the JD requirements, then use their facts.",
        "If a metric or technology is NOT in this library, do NOT invent it — use [X%] / [N users] placeholders.\n",
    ]
    for proj in projects:
        lines.append(f"▸ PROJECT: {proj['title']}")
        if proj["technologies"]:
            lines.append(f"  Technologies : {', '.join(proj['technologies'])}")
        if proj["metrics"]:
            lines.append(f"  Metrics      : {' | '.join(proj['metrics'])}")
        if proj["contributions"]:
            lines.append(f"  Contributions: {proj['contributions']}")
        if proj["challenges"]:
            lines.append(f"  Challenges   : {proj['challenges']}")
        if proj["summary"]:
            lines.append(f"  Summary      : {proj['summary']}")
        lines.append("")
    lines.append("━━━ END PROJECT LIBRARY ━━━")
    return "\n".join(lines)


# Per-job agents: Resume, Gap, Interview, Study are keyed by job_id. Coach & Partner are global.
PREP_AGENTS = ("resume", "gap", "interview", "study")
GLOBAL_AGENTS = ("coach", "partner", "synthesizer", "outreach")

def _migrate_conversations(raw: dict, jobs: list) -> dict:
    """Migrate old flat format to per-job format. Coach & Partner stay as lists."""
    out = {}
    out["coach"] = raw.get("coach", []) if isinstance(raw.get("coach"), list) else []
    out["partner"] = raw.get("partner", []) if isinstance(raw.get("partner"), list) else []
    out["synthesizer"] = raw.get("synthesizer", []) if isinstance(raw.get("synthesizer"), list) else []
    out["outreach"] = raw.get("outreach", []) if isinstance(raw.get("outreach"), list) else []

    def to_per_job(key: str) -> dict:
        val = raw.get(key, [])
        if isinstance(val, dict):
            return {int(k): v for k, v in val.items()}
        if not isinstance(val, list):
            val = []
        target_id = None
        jds_first = sorted(jobs, key=lambda j: (0 if j.get("jd") else 1, -j.get("created_at", 0)))
        if jds_first and jds_first[0].get("jd"):
            target_id = jds_first[0]["id"]
        elif jobs:
            target_id = jobs[0]["id"]
        if target_id is not None and val:
            return {str(target_id): val}
        return {}

    for k in PREP_AGENTS:
        out[k] = to_per_job(k)
    return out

# ─── STATE INIT ───
def init_state():
    # UI-only state (never persisted)
    ui_defaults = {
        "api_key": "",
        "current_page": "dashboard",
        "current_agent": "coach",
        "selected_job_id": None,
        "prep_job_id": None,  # Job selected for Resume/Gap/Interview/Study agents
    }
    for k, v in ui_defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

    # Persisted state — load from SQLite on first run this session
    if "_db_loaded" not in st.session_state:
        saved = db_load_all()
        st.session_state.profile = saved.get("profile", {
            "name": "", "role": "", "resume": "", "goals": "",
            "university": "", "gpa": "", "graduation": "",
            "resume_constraints": ""
        })
        # Back-fill new fields for existing saved profiles
        for _f in ("university", "gpa", "graduation", "resume_constraints"):
            if _f not in st.session_state.profile:
                st.session_state.profile[_f] = ""
        st.session_state.jobs = saved.get("jobs", [])
        st.session_state.next_job_id = saved.get("next_job_id", 1)
        raw_convos = saved.get("conversations", {})
        st.session_state.conversations = _migrate_conversations(raw_convos, saved.get("jobs", []))
        st.session_state._db_loaded = True

init_state()

def _normalize_api_key(raw_key: str) -> str:
    """Trim whitespace and a single layer of matching quotes."""
    key = (raw_key or "").strip()
    if len(key) >= 2 and key[0] == key[-1] and key[0] in {"'", '"'}:
        key = key[1:-1].strip()
    return key

def _read_env_api_key() -> str:
    key = _normalize_api_key(os.environ.get("ANTHROPIC_API_KEY", ""))
    if key:
        return key
    try:
        return _normalize_api_key(st.secrets.get("ANTHROPIC_API_KEY", ""))
    except Exception:
        return ""

def _looks_like_anthropic_key(key: str) -> bool:
    return key.startswith("sk-ant-")

DEFAULT_MAX_OUTPUT_TOKENS = 2400
MAX_CONTINUATIONS = 3
CONTINUE_PROMPT = "Continue exactly where you left off. Do not repeat earlier content. Finish the response."

def _extract_text_blocks(content_blocks) -> str:
    parts = []
    for block in content_blocks:
        if hasattr(block, "text") and block.text:
            parts.append(block.text)
    return "\n\n".join(parts).strip()

def _get_prep_messages(agent: str, job_id: int) -> list:
    convos = st.session_state.conversations
    if agent not in convos or not isinstance(convos[agent], dict):
        return []
    d = convos[agent]
    return d.get(str(job_id), d.get(job_id, []))  # JSON keys are strings

def _set_prep_messages(agent: str, job_id: int, messages: list) -> None:
    convos = st.session_state.conversations
    if agent not in convos or not isinstance(convos[agent], dict):
        convos[agent] = {}
    convos[agent][str(job_id)] = messages  # JSON keys are strings
    db_save("conversations", convos)

# ─── GOOGLE CALENDAR ───
GOOGLE_CALENDAR_SCOPES = ["https://www.googleapis.com/auth/calendar"]
CALENDAR_CREDS_KEY = "google_calendar_credentials"

def get_calendar_credentials():
    """Load stored Google Calendar credentials. Returns Credentials or None."""
    try:
        from google.oauth2.credentials import Credentials
        from google.auth.transport.requests import Request
        saved = db_load(CALENDAR_CREDS_KEY)
        if not saved or not isinstance(saved, dict):
            return None
        creds = Credentials.from_authorized_user_info(saved, GOOGLE_CALENDAR_SCOPES)
        if creds.expired and creds.refresh_token:
            creds.refresh(Request())
            db_save(CALENDAR_CREDS_KEY, json.loads(creds.to_json()))
        return creds if creds.valid else None
    except Exception:
        return None

def is_calendar_connected() -> bool:
    return get_calendar_credentials() is not None

def connect_calendar_oauth() -> tuple[bool, str]:
    """Run OAuth flow. Returns (success, message)."""
    try:
        from google_auth_oauthlib.flow import InstalledAppFlow
        from google.oauth2.credentials import Credentials
        creds_path = Path(__file__).parent / "google_credentials.json"
        if not creds_path.exists():
            return False, (
                "Missing google_credentials.json. "
                "Download OAuth credentials from Google Cloud Console (APIs & Services → Credentials → Create OAuth 2.0 Client ID → Desktop app) "
                "and save as google_credentials.json in the project folder."
            )
        flow = InstalledAppFlow.from_client_secrets_file(str(creds_path), GOOGLE_CALENDAR_SCOPES)
        creds = flow.run_local_server(port=0)
        db_save(CALENDAR_CREDS_KEY, json.loads(creds.to_json()))
        return True, "Google Calendar connected successfully."
    except Exception as e:
        return False, str(e)

def disconnect_calendar() -> None:
    con = _db()
    con.execute("DELETE FROM kv WHERE key = ?", (CALENDAR_CREDS_KEY,))
    con.commit()
    con.close()

def get_calendar_events(days_ahead: int = 7, max_results: int = 30) -> list[dict]:
    """Fetch upcoming events from primary calendar."""
    creds = get_calendar_credentials()
    if not creds:
        return []
    try:
        from googleapiclient.discovery import build
        from datetime import datetime, timezone, timedelta
        service = build("calendar", "v3", credentials=creds)
        now = datetime.now(timezone.utc).isoformat()
        time_max = (datetime.now(timezone.utc) + timedelta(days=days_ahead)).isoformat()
        events_result = (
            service.events()
            .list(
                calendarId="primary",
                timeMin=now,
                timeMax=time_max,
                maxResults=max_results,
                singleEvents=True,
                orderBy="startTime",
            )
            .execute()
        )
        events = events_result.get("items", [])
        out = []
        for e in events:
            start = e.get("start", {}) or {}
            start_str = start.get("dateTime") or start.get("date", "")
            out.append({"summary": e.get("summary", "(No title)"), "start": start_str})
        return out
    except Exception:
        return []

def create_calendar_event(title: str, start_datetime: str, end_datetime: str | None = None, description: str = "") -> tuple[bool, str]:
    """
    Create a calendar event. start_datetime and end_datetime in ISO format (YYYY-MM-DDTHH:MM or full).
    If end_datetime is None, default 1 hour.
    Returns (success, message).
    """
    creds = get_calendar_credentials()
    if not creds:
        return False, "Google Calendar not connected."
    try:
        from googleapiclient.discovery import build
        from datetime import datetime, timezone, timedelta
        service = build("calendar", "v3", credentials=creds)
        if not end_datetime:
            try:
                dt = datetime.fromisoformat(start_datetime.replace("Z", "+00:00"))
            except ValueError:
                dt = datetime.now(timezone.utc)
            end_dt = dt + timedelta(hours=1)
            end_datetime = end_dt.isoformat()
        if "T" not in start_datetime:
            start_datetime = start_datetime + "T09:00:00"
        if "T" not in end_datetime:
            end_datetime = end_datetime + "T10:00:00"
        if "+" not in start_datetime and not start_datetime.endswith("Z"):
            start_datetime += "Z"
        if "+" not in end_datetime and not end_datetime.endswith("Z"):
            end_datetime += "Z"
        body = {
            "summary": title,
            "start": {"dateTime": start_datetime, "timeZone": "UTC"},
            "end": {"dateTime": end_datetime, "timeZone": "UTC"},
        }
        if description:
            body["description"] = description
        service.events().insert(calendarId="primary", body=body).execute()
        return True, f"Created: {title}"
    except Exception as e:
        return False, str(e)

# Auto-load API key from environment or Streamlit secrets
env_key = _read_env_api_key()
if env_key and st.session_state.api_key != env_key:
    st.session_state.api_key = env_key

# ─── STATUS CONFIG ───
STATUSES = {
    "wishlist":  {"label": "Wishlist",      "color": "#7a7a8c", "icon": "◇", "badge": "badge-wishlist"},
    "applied":   {"label": "Applied",       "color": "#5bc8f5", "icon": "→", "badge": "badge-applied"},
    "screen":    {"label": "Phone Screen",  "color": "#f5c35b", "icon": "☎", "badge": "badge-screen"},
    "interview": {"label": "Interview",     "color": "#a55bf5", "icon": "✦", "badge": "badge-interview"},
    "offer":     {"label": "Offer",         "color": "#c5f135", "icon": "★", "badge": "badge-offer"},
    "rejected":  {"label": "Rejected",      "color": "#f55b7a", "icon": "✕", "badge": "badge-rejected"},
}

# ─── CLAUDE API ───
def get_client():
    key = _read_env_api_key() or _normalize_api_key(st.session_state.api_key)
    if key and st.session_state.api_key != key:
        st.session_state.api_key = key
    if not key:
        return None
    if not _looks_like_anthropic_key(key):
        return "invalid_format"
    return anthropic.Anthropic(api_key=key)

def call_claude(messages: list, system: str, max_tokens: int = DEFAULT_MAX_OUTPUT_TOKENS, model: str = "claude-opus-4-5") -> str:
    client = get_client()
    if not client:
        return "⚠️ No API key found. Set `ANTHROPIC_API_KEY` as an environment variable (e.g. `export ANTHROPIC_API_KEY=sk-ant-...`) and restart the app."
    if client == "invalid_format":
        return "⚠️ `ANTHROPIC_API_KEY` was loaded, but the value does not look like an Anthropic key. Update `.env` or your shell env and restart the app."
    try:
        msgs = list(messages)
        collected = []
        for _ in range(MAX_CONTINUATIONS + 1):
            response = client.messages.create(
                model=model,
                max_tokens=max_tokens,
                system=system,
                messages=msgs
            )
            text = _extract_text_blocks(response.content)
            if text:
                collected.append(text)
            if response.stop_reason != "max_tokens":
                return "\n\n".join(collected).strip()
            msgs.append({"role": "assistant", "content": response.content})
            msgs.append({"role": "user", "content": CONTINUE_PROMPT})
        return "\n\n".join(collected).strip()
    except anthropic.AuthenticationError:
        return "⚠️ Anthropic rejected the configured API key. Replace `ANTHROPIC_API_KEY` in `.env` (or your shell/Streamlit secret) with a current key and restart the app."
    except Exception as e:
        return f"⚠️ Error: {str(e)}"

_JUDGE_SYSTEM = """You are a resume quality judge evaluating rewritten resume bullets for an AI/ML/SWE internship candidate. Return ONLY a JSON object — no markdown, no prose."""

_JUDGE_PROMPT = """Evaluate the REWRITTEN BULLETS below against the JD and the candidate's source materials.

Score each dimension 0–10:
- jd_alignment: what % of the JD's key requirements are addressed by at least one bullet (10 = all key requirements covered)
- bullet_quality: what % of bullets follow the 2026 formula — strong verb + what + tech + metric or production signal (10 = all bullets ≥7/10)
- production_signal: do bullets show end-to-end ownership, latency, scale, deployment, or tradeoffs? (10 = strong production thinking throughout)
- grounding: are all claims traceable to the candidate's resume or projects, with no invented facts? (10 = fully grounded, 0 = heavy hallucination)

Also return:
- flags: list of up to 3 specific problems (each ≤15 words)
- strengths: list of up to 2 things working well (each ≤15 words)
- verdict: one sentence overall assessment

Return this exact JSON shape:
{{"jd_alignment": 0, "bullet_quality": 0, "production_signal": 0, "grounding": 0, "flags": [], "strengths": [], "verdict": ""}}

---
JD (first 2000 chars):
{jd}

CANDIDATE SOURCE MATERIALS (resume + projects, first 2000 chars):
{source}

REWRITTEN BULLETS TO EVALUATE:
{bullets}"""

def run_judge_scorecard(bullets: str, jd: str, source: str) -> dict:
    """Run a fast judge evaluation on rewritten resume bullets. Returns score dict or empty dict on failure."""
    import json as _j
    prompt = _JUDGE_PROMPT.format(
        jd=jd[:2000],
        source=source[:2000],
        bullets=bullets[:3000],
    )
    raw = call_claude(
        messages=[{"role": "user", "content": prompt}],
        system=_JUDGE_SYSTEM,
        max_tokens=512,
        model="claude-haiku-4-5-20251001",
    )
    try:
        cleaned = raw.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
        return _j.loads(cleaned)
    except Exception:
        return {}

def stream_claude(messages: list, system: str):
    client = get_client()
    if not client:
        yield "⚠️ No API key found. Set `ANTHROPIC_API_KEY` as an environment variable and restart the app."
        return
    if client == "invalid_format":
        yield "⚠️ `ANTHROPIC_API_KEY` was loaded, but the value does not look like an Anthropic key. Update `.env` or your shell env and restart the app."
        return
    try:
        with client.messages.stream(
            model="claude-opus-4-5",
            max_tokens=DEFAULT_MAX_OUTPUT_TOKENS,
            system=system,
            messages=messages
        ) as stream:
            for text in stream.text_stream:
                yield text
    except Exception as e:
        yield f"⚠️ Error: {str(e)}"

# ─── CALENDAR TOOLS (for Career Coach) ───
CALENDAR_TOOLS = [
    {
        "name": "get_calendar_events",
        "description": "Get the user's upcoming Google Calendar events. Use when planning their schedule, suggesting when to study, or checking availability. Returns event titles and start times.",
        "input_schema": {
            "type": "object",
            "properties": {
                "days_ahead": {"type": "integer", "description": "How many days ahead to fetch (default 7)", "default": 7},
            },
        },
    },
    {
        "name": "create_calendar_event",
        "description": "Create a new event on the user's Google Calendar. Use when the user wants to schedule study time, interview prep, application deadlines, or coaching follow-ups.",
        "input_schema": {
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Event title"},
                "start_datetime": {"type": "string", "description": "Start in ISO format: YYYY-MM-DDTHH:MM or YYYY-MM-DD"},
                "end_datetime": {"type": "string", "description": "End in same format. If omitted, 1 hour after start."},
                "description": {"type": "string", "description": "Optional event description"},
            },
            "required": ["title", "start_datetime"],
        },
    },
]

AGENT_TOOLS = [
    {
        "name": "run_gap_analysis",
        "description": "Run a full gap analysis for a specific job — fit score, critical gaps, closeable gaps, action plan. Use when user asks how well they match a role, whether to apply, or what gaps to close.",
        "input_schema": {
            "type": "object",
            "properties": {
                "job_id": {"type": "integer", "description": "Numeric job ID from the tracker (shown as [ID:N])."},
                "question": {"type": "string", "description": "Optional: specific focus (e.g. 'focus on the ML gap')."},
            },
            "required": ["job_id"],
        },
    },
    {
        "name": "run_resume_review",
        "description": "Rewrite and optimize resume bullets for a specific job — keyword alignment, ATS, copy-paste ready output. Use when user wants to tailor their resume to a role.",
        "input_schema": {
            "type": "object",
            "properties": {
                "job_id": {"type": "integer", "description": "Numeric job ID from the tracker."},
                "instruction": {"type": "string", "description": "Optional: e.g. 'show full analysis', 'rewrite projects only'."},
            },
            "required": ["job_id"],
        },
    },
    {
        "name": "draft_outreach",
        "description": "Draft a LinkedIn DM, connection request, referral ask, cold email, or follow-up for a target company. Use when user wants to reach out to someone.",
        "input_schema": {
            "type": "object",
            "properties": {
                "company": {"type": "string", "description": "Target company name."},
                "message_type": {
                    "type": "string",
                    "enum": ["linkedin_connection", "linkedin_dm", "referral_request", "cold_email", "follow_up"],
                    "description": "Type of outreach message.",
                },
                "contact_name": {"type": "string", "description": "Optional: name/role of the recipient."},
                "context": {"type": "string", "description": "Optional: how they found this person, shared connections, role to reference."},
            },
            "required": ["company", "message_type"],
        },
    },
    {
        "name": "run_interview_prep",
        "description": "Run mock interview or interview prep for a specific job. Use when user wants practice questions or prep for an upcoming interview.",
        "input_schema": {
            "type": "object",
            "properties": {
                "job_id": {"type": "integer", "description": "Numeric job ID from the tracker."},
                "focus": {"type": "string", "description": "Optional: 'behavioral', 'technical', 'system design', 'full round'."},
            },
            "required": ["job_id"],
        },
    },
    {
        "name": "run_study_plan",
        "description": "Build a prioritized study plan for a specific job or across all active applications. Use when user asks what to study or wants a prep schedule.",
        "input_schema": {
            "type": "object",
            "properties": {
                "job_id": {"type": "integer", "description": "Optional: job ID for a role-specific plan. Omit for cross-application plan."},
                "focus": {"type": "string", "description": "Optional: e.g. '2-week timeline', 'ML fundamentals only'."},
            },
            "required": [],
        },
    },
    {
        "name": "synthesize_resume",
        "description": "Analyze patterns across all tracked JDs and produce optimized multi-company resume versions. Use when user wants generalized resume strategy across multiple targets.",
        "input_schema": {
            "type": "object",
            "properties": {
                "instruction": {"type": "string", "description": "Optional: e.g. 'focus on ML/DS roles', 'what skills appear most'."},
            },
            "required": [],
        },
    },
]

def _run_calendar_tool(name: str, input_data: dict) -> str:
    if name == "get_calendar_events":
        days = input_data.get("days_ahead", 7)
        events = get_calendar_events(days_ahead=days)
        if not events:
            return "No upcoming events in that period."
        return "\n".join([f"- {e['start']}: {e['summary']}" for e in events])
    elif name == "create_calendar_event":
        title = input_data.get("title", "")
        start = input_data.get("start_datetime") or input_data.get("start")
        if not title or not start:
            return "Error: create_calendar_event requires 'title' and 'start_datetime' (e.g. 2025-03-21T14:00 or 2025-03-21)."
        ok, msg = create_calendar_event(
            title=title,
            start_datetime=start,
            end_datetime=input_data.get("end_datetime") or input_data.get("end"),
            description=input_data.get("description", ""),
        )
        return msg if ok else f"Error: {msg}"
    return "Unknown tool."

def _run_agent_tool(name: str, input_data: dict) -> str:
    """Dispatch a coach tool call to the appropriate sub-agent."""
    SUB_AGENT_MAX_TOKENS = 2000

    if name == "run_gap_analysis":
        job = get_job(input_data.get("job_id"))
        if not job:
            return f"Error: No job found with ID {input_data.get('job_id')}. Check tracker for valid IDs."
        q = input_data.get("question", "")
        msg = f"Run a full gap analysis for {job['company']} — {job['role']}." + (f" Focus: {q}" if q else "")
        return call_claude([{"role": "user", "content": msg}], get_system_prompt("gap", job=job), max_tokens=SUB_AGENT_MAX_TOKENS)

    elif name == "run_resume_review":
        job = get_job(input_data.get("job_id"))
        if not job:
            return f"Error: No job found with ID {input_data.get('job_id')}."
        ins = input_data.get("instruction", "")
        msg = f"Review and rewrite the resume for {job['company']} — {job['role']}." + (f" {ins}" if ins else "")
        return call_claude([{"role": "user", "content": msg}], get_system_prompt("resume", job=job), max_tokens=SUB_AGENT_MAX_TOKENS)

    elif name == "draft_outreach":
        company = input_data.get("company", "")
        mtype = input_data.get("message_type", "linkedin_dm").replace("_", " ")
        contact = input_data.get("contact_name", "")
        ctx = input_data.get("context", "")
        msg = f"Draft a {mtype} to {contact or 'someone'} at {company}." + (f" Context: {ctx}" if ctx else "")
        return call_claude([{"role": "user", "content": msg}], get_system_prompt("outreach"), max_tokens=SUB_AGENT_MAX_TOKENS)

    elif name == "run_interview_prep":
        job = get_job(input_data.get("job_id"))
        if not job:
            return f"Error: No job found with ID {input_data.get('job_id')}."
        focus = input_data.get("focus", "")
        msg = f"Run interview prep for {job['company']} — {job['role']}." + (f" Focus: {focus}" if focus else " Start with the most likely question types.")
        return call_claude([{"role": "user", "content": msg}], get_system_prompt("interview", job=job), max_tokens=SUB_AGENT_MAX_TOKENS)

    elif name == "run_study_plan":
        job_id = input_data.get("job_id")
        job = get_job(job_id) if job_id else None
        focus = input_data.get("focus", "")
        msg = (f"Build a study plan for {job['company']} — {job['role']}." if job else "Build a study plan across all my active applications.") + (f" Focus: {focus}" if focus else "")
        return call_claude([{"role": "user", "content": msg}], get_system_prompt("study", job=job), max_tokens=SUB_AGENT_MAX_TOKENS)

    elif name == "synthesize_resume":
        ins = input_data.get("instruction", "")
        msg = "Analyze patterns across all tracked JDs and produce optimized resume versions." + (f" {ins}" if ins else "")
        return call_claude([{"role": "user", "content": msg}], get_system_prompt("synthesizer"), max_tokens=SUB_AGENT_MAX_TOKENS)

    return f"Unknown agent tool: {name}"


def _run_tool(name: str, input_data: dict) -> str:
    """Route tool calls to calendar tools or agent tools."""
    if name in {t["name"] for t in CALENDAR_TOOLS}:
        return _run_calendar_tool(name, input_data)
    return _run_agent_tool(name, input_data)


def call_claude_with_tools(messages: list, system: str, tools: list, max_tokens: int = DEFAULT_MAX_OUTPUT_TOKENS) -> str:
    """Call Claude with tool use; executes tools and loops until done."""
    client = get_client()
    if not client:
        return "⚠️ No API key found. Set `ANTHROPIC_API_KEY` and restart."
    msgs = list(messages)
    max_rounds = 10
    continuation_count = 0
    collected = []
    for _ in range(max_rounds):
        response = client.messages.create(
            model="claude-opus-4-5",
            max_tokens=max_tokens,
            system=system,
            messages=msgs,
            tools=tools,
        )
        text = _extract_text_blocks(response.content)
        if text:
            collected.append(text)
        if response.stop_reason == "end_turn":
            return "\n\n".join(collected).strip()
        tool_results = []
        for block in response.content:
            if getattr(block, "type", None) == "tool_use":
                result = _run_tool(block.name, block.input)
                tool_results.append({"type": "tool_result", "tool_use_id": block.id, "content": result})
        if tool_results:
            msgs.append({"role": "assistant", "content": response.content})
            msgs.append({"role": "user", "content": tool_results})
            continue
        if response.stop_reason == "max_tokens" and continuation_count < MAX_CONTINUATIONS:
            continuation_count += 1
            msgs.append({"role": "assistant", "content": response.content})
            msgs.append({"role": "user", "content": CONTINUE_PROMPT})
            continue
        if collected:
            return "\n\n".join(collected).strip()
        msgs.append({"role": "assistant", "content": response.content})
    return "⚠️ Tool loop limit reached."

# ─── AGENT SYSTEM PROMPTS ───
from prompts import build_system_prompt as _build_system_prompt
from graph.project_matcher import build_matched_project_ctx as _match_project_ctx

def get_system_prompt(agent: str, job: dict | None = None) -> str:
    """Thin wrapper — delegates to prompts.build_system_prompt with session data."""
    p = st.session_state.profile
    jobs = st.session_state.jobs
    career_state = st.session_state.get("career_state", {})

    # ProjectMatcher: for resume agent, use top-3 matched projects instead of all
    if agent == "resume" and job and job.get("jd"):
        project_ctx = career_state.get("_matched_project_ctx") or build_project_context_for_resume()
    else:
        project_ctx = build_project_context_for_resume()

    resume_conversations = {}
    if agent == "synthesizer":
        raw_convos = st.session_state.conversations.get("resume", {})
        if isinstance(raw_convos, dict):
            resume_conversations = raw_convos

    return _build_system_prompt(
        agent=agent,
        profile=p,
        jobs=jobs,
        job=job,
        project_ctx=project_ctx,
        calendar_connected=is_calendar_connected(),
        resume_conversations=resume_conversations,
        weak_areas=career_state.get("weak_areas", []),
        gap_findings=career_state.get("gap_findings", ""),
    )


# ─── JOB HELPERS ───
def get_jobs():
    return st.session_state.jobs

def add_job(job_data: dict):
    job_data["id"] = st.session_state.next_job_id
    job_data["created_at"] = datetime.now().timestamp()
    st.session_state.jobs.append(job_data)
    st.session_state.next_job_id += 1
    db_save("jobs", st.session_state.jobs)
    db_save("next_job_id", st.session_state.next_job_id)

def update_job(job_id: int, updates: dict):
    for i, j in enumerate(st.session_state.jobs):
        if j["id"] == job_id:
            st.session_state.jobs[i].update(updates)
            break
    db_save("jobs", st.session_state.jobs)

def delete_job(job_id: int):
    st.session_state.jobs = [j for j in st.session_state.jobs if j["id"] != job_id]
    db_save("jobs", st.session_state.jobs)

def get_job(job_id: int) -> Optional[dict]:
    return next((j for j in st.session_state.jobs if j["id"] == job_id), None)

_VENV_PYTHON = str(Path(__file__).parent / ".venv" / "bin" / "python")

_FETCH_SCRIPT = """
import sys, json, re
url = sys.argv[1]
try:
    from playwright.sync_api import sync_playwright
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.goto(url, wait_until="domcontentloaded", timeout=20000)
        page.wait_for_timeout(1500)
        for selector in [
            "[class*='job-description']", "[class*='jobDescription']",
            "[class*='description']", "[id*='job-description']",
            "[id*='jobDescription']", "article", "main",
        ]:
            el = page.query_selector(selector)
            if el:
                text = el.inner_text()
                if len(text) > 200:
                    browser.close()
                    print(json.dumps({"text": re.sub(r"\\n{3,}", "\\n\\n", text.strip())[:8000], "error": ""}))
                    sys.exit(0)
        text = page.inner_text("body")
        browser.close()
        print(json.dumps({"text": re.sub(r"\\n{3,}", "\\n\\n", text.strip())[:8000], "error": ""}))
except Exception as e:
    print(json.dumps({"text": "", "error": str(e)}))
"""

def fetch_jd_from_url(url: str) -> tuple[str, str]:
    """Fetch and extract job description text from a URL using Playwright.
    Returns (jd_text, error_message). One will be empty string."""
    import subprocess, json as _json
    if not url or not url.startswith("http"):
        return "", "No valid URL found. Add a Job URL in the Details tab first."
    if "linkedin.com" in url:
        return "", "LinkedIn requires login — please paste the JD manually."
    try:
        result = subprocess.run(
            [_VENV_PYTHON, "-c", _FETCH_SCRIPT, url],
            capture_output=True, text=True, timeout=40
        )
        if result.returncode != 0 and not result.stdout.strip():
            return "", f"Subprocess error: {result.stderr.strip()[:300]}"
        data = _json.loads(result.stdout.strip())
        return data["text"], data["error"]
    except Exception as e:
        return "", f"Could not fetch page: {e}"

def autofill_job_from_url(url: str) -> tuple[dict, str]:
    """Fetch a job page and extract structured fields + JD using Claude.
    Returns (fields_dict, error_message). fields_dict keys: company, role, location, salary, jd."""
    raw_text, err = fetch_jd_from_url(url)
    if err:
        return {}, err
    prompt = f"""Extract job posting details from this text. Return ONLY a JSON object with these keys:
- company: company name (string)
- role: job title (string)
- location: city/state/remote (string, empty string if not found)
- salary: salary range (string, empty string if not found)
- jd: the full job description text, preserving responsibilities and requirements

Text:
{raw_text[:6000]}

Return only the JSON object, no markdown, no explanation."""
    import json as _json2
    try:
        response = call_claude(
            messages=[{"role": "user", "content": prompt}],
            system="You are a precise data extractor. Return only valid JSON.",
            max_tokens=4000,
        )
        # Strip markdown code fences if present
        cleaned = response.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
        fields = _json2.loads(cleaned)
        # jd fallback: use raw_text if Claude didn't populate it
        if not fields.get("jd") and raw_text:
            fields["jd"] = raw_text
        return fields, ""
    except Exception:
        # Claude call failed or JSON parse failed — still return the raw text as JD
        return {"jd": raw_text}, ""

def job_stats():
    jobs = get_jobs()
    total = len([j for j in jobs if j["status"] != "wishlist"])
    active = len([j for j in jobs if j["status"] in ("applied", "screen", "interview")])
    interviews = len([j for j in jobs if j["status"] in ("interview", "offer")])
    offers = len([j for j in jobs if j["status"] == "offer"])
    screened = len([j for j in jobs if j["status"] not in ("wishlist", "applied")])
    rate = f"{round(screened/total*100)}%" if total > 0 else "—"
    return total, active, interviews, offers, rate

AGENT_META = {
    "coach":       ("🧭", "Career Coach", "Monitors your pipeline · diagnoses blocks · builds strategy · holds you accountable"),
    "resume":      ("✏️", "Resume Reviewer", "10-sec scan test · bullet rewrites · ATS keywords · positioning strategy"),
    "gap":         ("🔍", "Gap Identifier", "Fit score · critical gaps · what to close before interviews"),
    "interview":   ("🎤", "Interview Coach", "Role-specific questions · realistic pressure · rubric-based feedback"),
    "study":       ("📚", "Study Planner", "Must-know topics · best resources · output-first learning schedule"),
    "partner":     ("🤝", "Study Partner", "Output-first · Feynman method · WHY before HOW · mistake patterns"),
    "synthesizer": ("⚡", "Resume Synthesizer", "Finds patterns across all your JDs · builds one great resume per role category · saves you from customizing every single application"),
    "outreach":    ("✉️", "LinkedIn & Outreach", "Cold DMs · referral requests · follow-ups · messages that actually get replies"),
}

DISPATCH_LABELS = {
    "skills_gap":     "🔍 Gap Analysis → 📚 Study Planner → 🧭 Synthesis",
    "execution":      "📂 Project Matcher → ✏️ Resume Review → 🧭 Synthesis",
    "interview_prep": "🎤 Interview Coach → 📚 Study Planner → 🧭 Synthesis",
    "strategy":       "⚡ Resume Synthesizer",
    "outreach":       "✉️ LinkedIn & Outreach",
}

STARTERS = {
    "coach":       ["Diagnose my pipeline. Where am I losing candidates — volume, response rate, or conversion?", "Give me 3 options for what to focus on this week — conservative, ambitious, and fastest ROI.", "Plan my week. I need time for study, applications, and interview prep. Block it on my calendar.", "Plan my week and block study time"],
    "resume":      ["Rewrite my bullets for this role — give me copy-paste ready output", "show full analysis", "Rewrite my weakest project bullets using my Project Library", "Rewrite my Skills section for this JD"],
    "gap":         ["Give me the full gap analysis with fit score for this role", "Be brutally honest — is this a realistic application or a reach? Score it and explain", "What are the dealbreaker gaps I need to address before I apply?", "What should I study or build in the next 2 weeks to close the most critical gap?"],
    "interview":   ["Run a mock technical interview", "Run a full behavioral round. Use STAR follow-ups if my answers are vague", "What patterns do you see in my answers so far? What's my biggest weakness?", "Test my system design knowledge"],
    "study":       ["Build a study plan for my active applications", "Build me a 4-week study plan for this role. Classify every topic as must-know, should-know, or good-to-know", "Create a 2-week prep schedule", "List must-know ML concepts"],
    "partner":     ["Explain transformers to me", "Quiz me on what I should know", "Teach me system design", "Help me understand RAG"],
    "synthesizer": ["Analyze patterns across all my JDs", "Build me a generalized SWE Intern resume", "What skills appear in most of my target JDs?", "Build one optimized resume version for [DS / SWE / AI-ML] roles across all my tracked companies"],
    "outreach":    ["Draft a LinkedIn connection request to a software engineer at [Company]", "I want to ask for a referral at [Company] — write a message for my 2nd touchpoint", "Write a cold DM to a recruiter at [Company] — I haven't applied yet", "Draft a follow-up message — I connected last week but no reply yet"],
}


def render_agent_chat(agent_key: str, *, show_job_selector: bool = True, chat_input_key: str = "", skip_chat_input: bool = False):
    """Render the full chat UI for a given agent. Reused by both the Agents page and Job Workbench.

    Args:
        agent_key: which agent to render (e.g. "coach", "resume")
        show_job_selector: whether to show the job dropdown for prep agents
        chat_input_key: unique suffix for st.chat_input key (needed when multiple on one page)
        skip_chat_input: if True, don't render st.chat_input (caller handles it separately)
    """
    icon, name, subtitle = AGENT_META[agent_key]
    is_prep_agent = agent_key in PREP_AGENTS

    p = st.session_state.profile
    jobs = get_jobs()
    jobs_with_jd = [j for j in jobs if j.get("jd")]

    # For prep agents: job selector + auto-default prep_job_id
    if is_prep_agent:
        if not jobs_with_jd:
            st.info("Add jobs and paste job descriptions in the **Job Tracker** (Notes & JD tab) to unlock Resume, Gap, Interview, and Study agents. Each job gets its own set of preparation agents.")
            st.stop()
        valid_prep_ids = [j["id"] for j in jobs_with_jd]
        if st.session_state.prep_job_id not in valid_prep_ids:
            st.session_state.prep_job_id = jobs_with_jd[0]["id"]
        prep_job = get_job(st.session_state.prep_job_id) or jobs_with_jd[0]
        if show_job_selector:
            job_options = [(j["id"], f"{j['company']} — {j['role']}") for j in jobs_with_jd]
            option_ids = [jid for jid, _ in job_options]
            option_labels = [lb for _, lb in job_options]
            curr_idx = option_ids.index(st.session_state.prep_job_id) if st.session_state.prep_job_id in option_ids else 0
            st.markdown(f'<div style="font-size:15px;font-weight:700;margin:0 0 8px 0;color:#eeedf0">{icon} {name} <span style="font-weight:400;color:#7a7a8c;font-size:12px">{subtitle}</span></div>', unsafe_allow_html=True)
            chosen_label = st.selectbox(
                "**Preparing for**",
                options=option_labels,
                index=curr_idx,
                key=f"prep_job_select_{chat_input_key or agent_key}",
            )
            chosen_id = option_ids[option_labels.index(chosen_label)]
            if chosen_id != st.session_state.prep_job_id:
                st.session_state.prep_job_id = chosen_id
                st.rerun()
        else:
            prep_job_label = f"{prep_job.get('company', '—')} — {prep_job.get('role', '—')}"
            st.markdown(f'<div style="font-size:15px;font-weight:700;margin:0 0 2px 0;color:#eeedf0">{icon} {name} <span style="font-weight:400;color:#7a7a8c;font-size:12px">· {prep_job_label}</span></div>', unsafe_allow_html=True)
    else:
        st.markdown(f'<div style="font-size:15px;font-weight:700;margin:0 0 2px 0;color:#eeedf0">{icon} {name} <span style="font-weight:400;color:#7a7a8c;font-size:12px">{subtitle}</span></div>', unsafe_allow_html=True)

    col_h, col_clear = st.columns([5, 1])
    with col_clear:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("Clear chat", key=f"clear_{chat_input_key or agent_key}"):
            if is_prep_agent and st.session_state.prep_job_id:
                _set_prep_messages(agent_key, st.session_state.prep_job_id, [])
            else:
                st.session_state.conversations[agent_key] = []
                db_save("conversations", st.session_state.conversations)
            st.rerun()

    # Context status bar
    has_resume = bool(p.get("resume"))
    has_jobs = len(jobs) > 0
    has_jd = len(jobs_with_jd) > 0
    ctx_parts = []
    ctx_parts.append(f"{'✅' if has_resume else '⚠️'} Resume {'set' if has_resume else 'missing — go to Profile'}")
    ctx_parts.append(f"{'✅' if has_jobs else '⚠️'} {len(jobs)} job{'s' if len(jobs) != 1 else ''} tracked")
    ctx_parts.append(f"{'✅' if has_jd else '⚠️'} {len(jobs_with_jd)} job{'s' if len(jobs_with_jd) != 1 else ''} with JD")

    st.markdown(f"""<div style="background:#101115;border:1px solid rgba(255,255,255,0.06);border-radius:8px;padding:8px 14px;font-family:'JetBrains Mono',monospace;font-size:12px;color:#7a7a8c;margin-bottom:16px;">
    {"&nbsp;&nbsp;·&nbsp;&nbsp;".join(ctx_parts)}
    </div>""", unsafe_allow_html=True)

    if is_prep_agent:
        messages = _get_prep_messages(agent_key, st.session_state.prep_job_id)
        prep_job = get_job(st.session_state.prep_job_id)
    else:
        messages = st.session_state.conversations.get(agent_key, [])

    if not messages:
        # Starter chips
        st.markdown('<div style="color:#7a7a8c;font-size:12px;margin:12px 0 8px">Try one of these →</div>', unsafe_allow_html=True)
        cols = st.columns(2)
        for i, s in enumerate(STARTERS.get(agent_key, [])):
            with cols[i % 2]:
                if st.button(f"↗ {s}", key=f"starter_{chat_input_key or agent_key}_{i}", use_container_width=True):
                    if is_prep_agent:
                        msgs = _get_prep_messages(agent_key, st.session_state.prep_job_id)
                        msgs.append({"role": "user", "content": s})
                        _set_prep_messages(agent_key, st.session_state.prep_job_id, msgs)
                    else:
                        st.session_state.conversations[agent_key].append({"role": "user", "content": s})
                        db_save("conversations", st.session_state.conversations)
                    st.rerun()
    else:
        # Render messages
        for msg in messages:
            if msg["role"] == "user":
                st.markdown(f'<div class="chat-label" style="text-align:right">YOU</div>', unsafe_allow_html=True)
                st.markdown(f'<div class="chat-msg-user" style="white-space:pre-wrap">{msg["content"]}</div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="chat-label agent-label">{icon} {name}</div>', unsafe_allow_html=True)
                # Check for dispatch indicator stored in the message
                if msg.get("dispatch"):
                    label = DISPATCH_LABELS.get(msg["dispatch"], "")
                    if label:
                        st.markdown(f'<div class="dispatch-indicator">⚙ {label}</div>', unsafe_allow_html=True)
                with st.container():
                    st.markdown(msg["content"])

        # Copy-ready panel + judge scorecard — Resume Reviewer only
        if agent_key == "resume":
            last_assistant = next(
                (m["content"] for m in reversed(messages) if m["role"] == "assistant"),
                None
            )
            if last_assistant:
                with st.expander("📋 Copy-ready output — paste directly into Google Docs", expanded=True):
                    st.text_area(
                        label="copy_box",
                        value=last_assistant,
                        height=220,
                        key=f"resume_copy_box_{chat_input_key or agent_key}",
                        label_visibility="collapsed"
                    )
                    st.caption("Select all (Ctrl+A / Cmd+A) → Copy → Paste into Docs. Plain text — no HTML artifacts.")

                # ── Judge scorecard (before / after) ──
                _job = get_job(st.session_state.prep_job_id) if st.session_state.prep_job_id else {}
                _jd  = _job.get("jd", "")
                _source = (p.get("resume") or "") + "\n\n" + build_project_context_for_resume()

                after_key  = str(hash(last_assistant + str(st.session_state.prep_job_id)))
                before_key = str(hash(_source + _jd + "before"))
                scorecard_cache = st.session_state.setdefault("resume_scorecard_cache", {})

                col_sc, col_reeval = st.columns([6, 1])
                with col_reeval:
                    if st.button("↺ Re-evaluate", key=f"reeval_scorecard_{chat_input_key or agent_key}", use_container_width=True):
                        scorecard_cache.pop(after_key, None)
                        scorecard_cache.pop(before_key, None)
                        st.rerun()

                if after_key not in scorecard_cache or before_key not in scorecard_cache:
                    with st.spinner("Evaluating before & after..."):
                        if before_key not in scorecard_cache:
                            scorecard_cache[before_key] = run_judge_scorecard(
                                p.get("resume") or "", _jd, _source
                            )
                        if after_key not in scorecard_cache:
                            scorecard_cache[after_key] = run_judge_scorecard(
                                last_assistant, _jd, _source
                            )

                sc_before = scorecard_cache.get(before_key, {})
                sc_after  = scorecard_cache.get(after_key, {})

                if sc_before and sc_after:
                    def _score_color(v):
                        if v >= 8: return "#4ade80"
                        if v >= 6: return "#facc15"
                        return "#f87171"

                    def _delta_html(b, a):
                        d = a - b
                        if d > 0:  return f'<span style="color:#4ade80">▲{d}</span>'
                        if d < 0:  return f'<span style="color:#f87171">▼{abs(d)}</span>'
                        return '<span style="color:#7a7a8c">—</span>'

                    dims = [
                        ("JD Alignment",      "jd_alignment"),
                        ("Bullet Quality",    "bullet_quality"),
                        ("Production Signal", "production_signal"),
                        ("Grounding",         "grounding"),
                    ]

                    rows = "".join(
                        f'<tr>'
                        f'<td style="padding:5px 12px 5px 0;color:#aaa">{label}</td>'
                        f'<td style="padding:5px 8px;color:{_score_color(sc_before.get(key,0))};font-family:JetBrains Mono,monospace;text-align:center">{sc_before.get(key,0)}/10</td>'
                        f'<td style="padding:5px 8px;color:{_score_color(sc_after.get(key,0))};font-family:JetBrains Mono,monospace;text-align:center">{sc_after.get(key,0)}/10</td>'
                        f'<td style="padding:5px 8px;font-family:JetBrains Mono,monospace;text-align:center">{_delta_html(sc_before.get(key,0), sc_after.get(key,0))}</td>'
                        f'</tr>'
                        for label, key in dims
                    )

                    st.markdown(
                        f'<table style="font-size:13px;border-collapse:collapse;margin:10px 0 8px">'
                        f'<thead><tr>'
                        f'<th style="padding:4px 12px 4px 0;color:#555568;font-weight:500;text-align:left"></th>'
                        f'<th style="padding:4px 8px;color:#555568;font-weight:500;text-align:center">Before</th>'
                        f'<th style="padding:4px 8px;color:#555568;font-weight:500;text-align:center">After</th>'
                        f'<th style="padding:4px 8px;color:#555568;font-weight:500;text-align:center">Δ</th>'
                        f'</tr></thead><tbody>{rows}</tbody></table>',
                        unsafe_allow_html=True,
                    )

                    if sc_after.get("verdict"):
                        st.markdown(f'<div style="font-size:13px;color:#aaa;margin-bottom:6px;">💬 {sc_after["verdict"]}</div>', unsafe_allow_html=True)

                    flag_lines     = "".join(f"<li>{f}</li>" for f in sc_after.get("flags", []))
                    strength_lines = "".join(f"<li>{s}</li>" for s in sc_after.get("strengths", []))
                    if flag_lines or strength_lines:
                        st.markdown(
                            '<div style="font-size:13px;color:#7a7a8c;line-height:1.7">'
                            + (f'<span style="color:#f87171">⚑ Still needs work:</span><ul style="margin:2px 0 6px 16px">{flag_lines}</ul>' if flag_lines else "")
                            + (f'<span style="color:#4ade80">✓ Working well:</span><ul style="margin:2px 0 0 16px">{strength_lines}</ul>' if strength_lines else "")
                            + '</div>',
                            unsafe_allow_html=True,
                        )

        # Scroll to the bottom so the latest message is always in view
        _st_components.html(
            "<script>setTimeout(function(){var m=window.parent.document.querySelector('[data-testid=\"stAppViewContainer\"] > section:first-child');if(m)m.scrollTop=999999;},120);</script>",
            height=0,
        )

        # If last message is from user, generate response
        if messages and messages[-1]["role"] == "user":
            prep_job = get_job(st.session_state.prep_job_id) if is_prep_agent else None
            api_messages = [{"role": m["role"], "content": m["content"]} for m in messages]
            career_state = st.session_state.setdefault("career_state", {})

            # ProjectMatcher: rank projects by JD relevance before resume system prompt
            if agent_key == "resume" and prep_job and prep_job.get("jd"):
                from graph.project_matcher import build_matched_project_ctx
                career_state["_matched_project_ctx"] = build_matched_project_ctx(
                    prep_job["jd"], db_get_projects()
                )

            system = get_system_prompt(agent_key, job=prep_job)

            with st.spinner(f"{name} is thinking..."):
                if agent_key == "coach":
                    from graph.graph import invoke_coach_graph
                    response, updated_career_state, diagnosis = invoke_coach_graph(
                        api_messages=api_messages,
                        profile=st.session_state.profile,
                        jobs=st.session_state.jobs,
                        project_ctx=build_project_context_for_resume(),
                        calendar_connected=is_calendar_connected(),
                        resume_conversations=st.session_state.conversations.get("resume", {}),
                        career_state=career_state,
                        thread_id=f"careeros_coach",
                    )
                    st.session_state.career_state = updated_career_state
                    st.session_state._last_coach_diagnosis = diagnosis
                else:
                    response = call_claude(api_messages, system)
                    # Cross-agent: extract weak areas after interview sessions
                    if agent_key == "interview" and response:
                        from graph.nodes import extract_weak_areas
                        found = extract_weak_areas(response)
                        if found:
                            career_state["weak_areas"] = found

            # Clear the per-session matched project ctx after use
            career_state.pop("_matched_project_ctx", None)

            # Build the assistant message dict (with optional dispatch metadata for coach)
            assistant_msg = {"role": "assistant", "content": response}
            if agent_key == "coach" and st.session_state.get("_last_coach_diagnosis"):
                diag = st.session_state._last_coach_diagnosis
                if diag != "direct":
                    assistant_msg["dispatch"] = diag

            if is_prep_agent:
                msgs = _get_prep_messages(agent_key, st.session_state.prep_job_id)
                msgs.append(assistant_msg)
                _set_prep_messages(agent_key, st.session_state.prep_job_id, msgs)
            else:
                st.session_state.conversations[agent_key].append(assistant_msg)
                db_save("conversations", st.session_state.conversations)
            st.rerun()

    # Input — sticky at bottom of viewport, always visible without scrolling
    if not skip_chat_input:
        user_input = st.chat_input(f"Ask {name}...", key=f"chat_input_{chat_input_key or agent_key}")
        if user_input and user_input.strip():
            if is_prep_agent:
                msgs = _get_prep_messages(agent_key, st.session_state.prep_job_id)
                msgs.append({"role": "user", "content": user_input.strip()})
                _set_prep_messages(agent_key, st.session_state.prep_job_id, msgs)
            else:
                st.session_state.conversations[agent_key].append({"role": "user", "content": user_input.strip()})
                db_save("conversations", st.session_state.conversations)
            st.rerun()


# ─── SIDEBAR ───
with st.sidebar:
    st.markdown('<div class="logo-text">CareerOS</div>', unsafe_allow_html=True)
    st.markdown('<div class="logo-sub">Job Hunt Command Center</div>', unsafe_allow_html=True)
    st.divider()

    # Navigation — streamlined 6-item sidebar
    if st.button("◈  Dashboard", use_container_width=True):
        st.session_state.current_page = "dashboard"
        st.rerun()
    if st.button("⊞  Job Tracker", use_container_width=True):
        st.session_state.current_page = "tracker"
        st.rerun()

    st.markdown('<div class="sidebar-section">AI</div>', unsafe_allow_html=True)

    # Career Coach — primary entry point, visually emphasized via CSS class
    st.markdown('<div class="coach-btn-wrap">', unsafe_allow_html=True)
    if st.button("🧭  Career Coach", use_container_width=True, key="nav_coach"):
        st.session_state.current_page = "agents"
        st.session_state.current_agent = "coach"
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

    if st.button("📋  Job Workbench", use_container_width=True, key="nav_workbench"):
        st.session_state.current_page = "workbench"
        st.rerun()

    st.markdown('<div class="sidebar-section">My Work</div>', unsafe_allow_html=True)
    proj_count = len(db_get_projects())
    proj_label = f"📚  Project Library ({proj_count})" if proj_count else "📚  Project Library"
    if st.button(proj_label, use_container_width=True):
        st.session_state.current_page = "project_library"
        st.rerun()

    st.divider()
    if st.button("⚙  My Profile", use_container_width=True):
        st.session_state.current_page = "profile"
        st.rerun()

    # Mini stats
    st.divider()
    total, active, interviews, offers, rate = job_stats()
    st.markdown(f"""
    <div style="font-family:'JetBrains Mono',monospace;font-size:12px;color:#7a7a8c;line-height:2;">
    📋 {total} applied &nbsp;·&nbsp; {active} active<br>
    🎤 {interviews} interviews &nbsp;·&nbsp; ★ {offers} offers<br>
    📊 {rate} response rate
    </div>
    """, unsafe_allow_html=True)

# ─── DASHBOARD ───
if st.session_state.current_page == "dashboard":
    st.title("Dashboard")
    total, active, interviews, offers, rate = job_stats()

    col1, col2, col3, col4, col5 = st.columns(5)
    with col1: st.metric("Total Applied", total)
    with col2: st.metric("Active Pipeline", active)
    with col3: st.metric("Interviews", interviews)
    with col4: st.metric("Offers", offers)
    with col5: st.metric("Response Rate", rate)

    st.divider()
    col_l, col_r = st.columns([3, 2])

    with col_l:
        st.subheader("Recent Applications")
        recent = sorted(get_jobs(), key=lambda x: x.get("created_at", 0), reverse=True)[:6]
        if not recent:
            st.info("No applications yet. Add your first job in the tracker!")
        else:
            for j in recent:
                s = STATUSES[j["status"]]
                col_a, col_b, col_c = st.columns([4, 1, 1])
                with col_a:
                    if st.button(f"**{j['company']}** — {j.get('role', '—')}", key=f"dash_job_{j['id']}", use_container_width=True):
                        st.session_state.current_page = "add_job"
                        st.session_state.editing_job_id = j["id"]
                        st.rerun()
                with col_b:
                    st.markdown(f"""<span class="badge badge-{j['status']}">{s['icon']} {s['label']}</span>""",
                               unsafe_allow_html=True)
                with col_c:
                    if j.get("jd"):
                        if st.button("Prep", key=f"dash_prep_{j['id']}"):
                            st.session_state.prep_job_id = j["id"]
                            st.session_state.current_page = "workbench"
                            st.rerun()

        if st.button("View All →", key="dash_view_all"):
            st.session_state.current_page = "tracker"
            st.rerun()

    with col_r:
        st.subheader("🧭 Coach Insight")
        jobs = get_jobs()
        profile = st.session_state.profile

        if not profile.get("resume") or not jobs:
            st.info("Fill in your profile and add some job applications to get a personalized coach insight.")
        else:
            col_btn, col_refresh = st.columns([3, 1])
            with col_btn:
                gen_pressed = st.button("Generate Insight", key="gen_insight")
            with col_refresh:
                if "dashboard_insight" in st.session_state and st.session_state.dashboard_insight:
                    if st.button("↺ Refresh", key="refresh_insight"):
                        st.session_state.dashboard_insight = ""
                        st.rerun()

            if gen_pressed or (not st.session_state.get("dashboard_insight") and False):
                # Check for approaching deadlines
                urgent = [j for j in jobs if j.get("deadline") and j["status"] not in ("offer", "rejected")]
                urgent_lines = []
                for j in urgent:
                    try:
                        days_left = (date.fromisoformat(j["deadline"]) - date.today()).days
                        if 0 <= days_left <= 14:
                            urgent_lines.append(f"- {j['company']} ({j['role']}): {days_left}d left")
                    except ValueError:
                        pass
                deadline_note = f"\nURGENT DEADLINES:\n" + "\n".join(urgent_lines) if urgent_lines else ""
                jobs_summary = "\n".join([
                    f"- {j['company']} ({j['role']}) — {STATUSES[j['status']]['label']}"
                    + (f" · deadline {j['deadline']}" if j.get('deadline') else "")
                    for j in jobs
                ])
                prompt = f"""Candidate: {profile['name'] or 'Annie'}, targeting: {profile['role'] or 'tech internship'}
University: {profile.get('university') or 'not set'} · Graduation: {profile.get('graduation') or 'not set'} · GPA: {profile.get('gpa') or 'not set'}
Job applications:
{jobs_summary}
{deadline_note}

Give ONE sharp, specific internship coach insight (3-4 sentences). Spot a pattern, risk, opportunity, or urgent deadline. Reference actual companies/roles. Flag if any deadlines are approaching."""

                with st.spinner("Analyzing your pipeline..."):
                    insight = call_claude([{"role": "user", "content": prompt}],
                                        "You are a top-tier internship career strategist. Be concise, specific, and actionable.")
                st.session_state.dashboard_insight = insight

            if st.session_state.get("dashboard_insight"):
                st.markdown(f"""
                <div style="background:#17191f;border:1px solid rgba(124,106,247,0.25);border-radius:10px;padding:16px;font-size:15px;line-height:1.7;color:#eeedf0;">
                🧭 {st.session_state.dashboard_insight}
                </div>""", unsafe_allow_html=True)

# ─── JOB TRACKER ───
elif st.session_state.current_page == "tracker":
    col_title, col_btn = st.columns([6, 1])
    with col_title:
        st.title("Job Tracker")
    with col_btn:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("+ Add Job", key="add_job_top"):
            st.session_state.current_page = "add_job"
            st.session_state.editing_job_id = None
            st.rerun()

    # Kanban
    tabs = st.tabs([f"{STATUSES[s]['icon']} {STATUSES[s]['label']}" for s in STATUSES])
    for i, status_key in enumerate(STATUSES):
        with tabs[i]:
            status_jobs = [j for j in get_jobs() if j["status"] == status_key]
            s = STATUSES[status_key]
            st.markdown(f"**{len(status_jobs)} application{'s' if len(status_jobs) != 1 else ''}**")

            if not status_jobs:
                st.markdown(f"""<div style="border:1px dashed rgba(255,255,255,0.1);border-radius:10px;padding:24px;text-align:center;color:#555568;font-size:14px;">No applications here yet</div>""",
                           unsafe_allow_html=True)
            else:
                for j in status_jobs:
                    with st.container():
                        c1, c2, c3, c4 = st.columns([5, 2, 1, 1])
                        with c1:
                            st.markdown(f"**{j['company']}** — {j.get('role', '—')}")
                            if j.get('location'):
                                st.caption(f"📍 {j['location']}" + (f" · {j['salary']}" if j.get('salary') else ""))
                            if j.get('date'):
                                st.caption(f"Applied: {j['date']}")
                            if j.get('deadline'):
                                try:
                                    dl = date.fromisoformat(j['deadline'])
                                    days_left = (dl - date.today()).days
                                    if days_left < 0:
                                        dl_color = "#f55b7a"
                                        dl_label = f"⚠️ Deadline passed ({j['deadline']})"
                                    elif days_left <= 7:
                                        dl_color = "#f5c35b"
                                        dl_label = f"⏳ {days_left}d left — {j['deadline']}"
                                    elif days_left <= 14:
                                        dl_color = "#f5c35b"
                                        dl_label = f"Deadline: {j['deadline']} ({days_left}d)"
                                    else:
                                        dl_color = "#7a7a8c"
                                        dl_label = f"Deadline: {j['deadline']}"
                                    st.markdown(f'<span style="font-size:14px;color:{dl_color}">{dl_label}</span>', unsafe_allow_html=True)
                                except ValueError:
                                    st.caption(f"Deadline: {j['deadline']}")
                        with c2:
                            if j.get('tags'):
                                st.caption(" · ".join(j['tags'][:3]))
                        with c3:
                            if st.button("Edit", key=f"edit_{j['id']}"):
                                st.session_state.current_page = "add_job"
                                st.session_state.editing_job_id = j["id"]
                                st.rerun()
                        with c4:
                            if j.get("jd"):
                                if st.button("Prep", key=f"prep_{j['id']}"):
                                    st.session_state.prep_job_id = j["id"]
                                    st.session_state.current_page = "workbench"
                                    st.rerun()
                        st.divider()

# ─── ADD / EDIT JOB ───
elif st.session_state.current_page == "add_job":
    editing_id = getattr(st.session_state, "editing_job_id", None)
    existing = get_job(editing_id) if editing_id else None

    st.title("Edit Job" if existing else "Add Job")

    tab1, tab3 = st.tabs(["📋 Details", "📄 Notes & JD"])

    with tab1:
        # ── Auto-fill from URL ──
        prefill = st.session_state.get("job_prefill", {})

        col_url_input, col_fetch_btn = st.columns([3, 1])
        with col_url_input:
            job_url = st.text_input("Job URL", value=prefill.get("url", existing.get("url", "") if existing else ""), key="form_job_url")
        with col_fetch_btn:
            st.markdown("<br>", unsafe_allow_html=True)
            fetch_clicked = st.button("🔗 Auto-fill from URL", use_container_width=True, disabled=not bool(job_url))

        if fetch_clicked:
            with st.spinner("Fetching job details..."):
                fields, err = autofill_job_from_url(job_url)
            if err:
                st.warning(err)
            else:
                fields["url"] = job_url
                st.session_state.job_prefill = fields
                if existing:
                    update_job(editing_id, {
                        "company": fields.get("company", existing.get("company", "")),
                        "role":    fields.get("role",    existing.get("role", "")),
                        "location": fields.get("location", existing.get("location", "")),
                        "salary":  fields.get("salary",  existing.get("salary", "")),
                        "url":     job_url,
                        "jd":      fields.get("jd",      existing.get("jd", "")),
                    })
                st.success("Done! Review the fields below and save.")
                st.rerun()

        # ── Form fields (use prefill values if present, else existing/blank) ──
        col1, col2 = st.columns(2)
        with col1:
            company = st.text_input("Company *", value=prefill.get("company", existing.get("company", "") if existing else ""))
            status = st.selectbox("Status", list(STATUSES.keys()),
                                  format_func=lambda x: f"{STATUSES[x]['icon']} {STATUSES[x]['label']}",
                                  index=list(STATUSES.keys()).index(existing.get("status", "applied")) if existing else 1)
            location = st.text_input("Location", value=prefill.get("location", existing.get("location", "") if existing else ""))

        with col2:
            role = st.text_input("Role / Title", value=prefill.get("role", existing.get("role", "") if existing else ""))
            applied_date = st.date_input("Date Applied",
                                         value=date.fromisoformat(existing["date"]) if existing and existing.get("date") else date.today())
            salary = st.text_input("Salary Range", value=prefill.get("salary", existing.get("salary", "") if existing else ""))

        col_deadline, _ = st.columns(2)
        with col_deadline:
            deadline_val = existing.get("deadline", "") if existing else ""
            deadline_input = st.text_input(
                "Application Deadline",
                value=deadline_val,
                placeholder="e.g. 2025-10-31",
                help="Keeps the Career Coach aware of urgency. Use YYYY-MM-DD format."
            )
        tags_input = st.text_input("Tags (comma separated)", value=", ".join(existing.get("tags", [])) if existing else "")

        col_save, col_del, col_cancel = st.columns([2, 1, 1])
        with col_save:
            if st.button("💾 Save Job", use_container_width=True):
                if not company:
                    st.error("Company name is required.")
                else:
                    job_data = {
                        "company": company, "role": role,
                        "status": status, "date": applied_date.isoformat(),
                        "location": location, "salary": salary,
                        "url": job_url,
                        "deadline": deadline_input.strip(),
                        "tags": [t.strip() for t in tags_input.split(",") if t.strip()],
                        "jd": prefill.get("jd", existing.get("jd", "") if existing else ""),
                        "notes": existing.get("notes", "") if existing else "",
                    }
                    if existing:
                        job_data["id"] = editing_id
                        job_data["created_at"] = existing.get("created_at", datetime.now().timestamp())
                        update_job(editing_id, job_data)
                        st.success("✓ Job updated!")
                    else:
                        add_job(job_data)
                        st.success("✓ Job added!")
                    st.session_state.pop("job_prefill", None)
                    st.session_state.current_page = "tracker"
                    st.rerun()

        with col_del:
            if existing and st.button("🗑 Delete", use_container_width=True):
                delete_job(editing_id)
                st.session_state.pop("job_prefill", None)
                st.session_state.current_page = "tracker"
                st.rerun()

        with col_cancel:
            if st.button("Cancel", use_container_width=True):
                st.session_state.pop("job_prefill", None)
                st.session_state.current_page = "tracker"
                st.rerun()

    with tab3:
        if not existing:
            st.info("Save the job first to add notes and JD.")
        else:
            jd_text = st.text_area("Job Description (paste here to power AI agents)",
                                   value=existing.get("jd", ""),
                                   height=250,
                                   placeholder="Paste the full job description...")
            notes_text = st.text_area("Personal Notes",
                                      value=existing.get("notes", ""),
                                      height=150,
                                      placeholder="Why you're interested, who referred you, interview impressions...")
            if st.button("Save Notes & JD"):
                update_job(editing_id, {"jd": jd_text, "notes": notes_text})
                st.success("Saved!")
                st.rerun()

# ─── AGENTS (Coach + global specialist agents) ───
elif st.session_state.current_page == "agents":
    agent_key = st.session_state.current_agent
    render_agent_chat(agent_key)

    # Specialist tools expander — only on Coach page
    if agent_key == "coach":
        with st.expander("Specialist Tools", expanded=False):
            tool_cols = st.columns(3)
            for i, (key, icon, label) in enumerate([
                ("partner", "🤝", "Study Partner"),
                ("synthesizer", "⚡", "Resume Synthesizer"),
                ("outreach", "✉️", "LinkedIn & Outreach"),
            ]):
                with tool_cols[i]:
                    if st.button(f"{icon} {label}", key=f"specialist_{key}", use_container_width=True):
                        st.session_state.current_agent = key
                        st.rerun()

# ─── JOB WORKBENCH (tabbed prep agents per job) ───
elif st.session_state.current_page == "workbench":
    jobs_with_jd = [j for j in get_jobs() if j.get("jd")]

    if not jobs_with_jd:
        st.title("Job Workbench")
        st.info("Add jobs and paste job descriptions in the **Job Tracker** (Notes & JD tab) to unlock the Job Workbench.")
        st.stop()

    # Job selector at top
    valid_prep_ids = [j["id"] for j in jobs_with_jd]
    if st.session_state.prep_job_id not in valid_prep_ids:
        st.session_state.prep_job_id = jobs_with_jd[0]["id"]

    job_options = [(j["id"], f"{j['company']} — {j['role']}") for j in jobs_with_jd]
    option_ids = [jid for jid, _ in job_options]
    option_labels = [lb for _, lb in job_options]
    curr_idx = option_ids.index(st.session_state.prep_job_id) if st.session_state.prep_job_id in option_ids else 0

    # ── Job selector + agent tabs in the sidebar (always visible, no scrolling needed) ──
    with st.sidebar:
        st.markdown('<div class="sidebar-section">Workbench</div>', unsafe_allow_html=True)
        chosen_label = st.selectbox(
            "**Preparing for**",
            options=option_labels,
            index=curr_idx,
            key="workbench_job_select",
        )
        chosen_id = option_ids[option_labels.index(chosen_label)]
        if chosen_id != st.session_state.prep_job_id:
            st.session_state.prep_job_id = chosen_id
            st.rerun()

        # Agent tabs as sidebar buttons
        if "workbench_tab" not in st.session_state:
            st.session_state.workbench_tab = "resume"

        tab_labels = ["✏️ Resume Reviewer", "🔍 Gap Identifier", "🎤 Interview Coach", "📚 Study Planner"]
        tab_keys = ["resume", "gap", "interview", "study"]

        for key, label in zip(tab_keys, tab_labels):
            is_active = st.session_state.workbench_tab == key
            btn_class = "workbench-tab-active" if is_active else ""
            st.markdown(f'<div class="{btn_class}">', unsafe_allow_html=True)
            if st.button(label, key=f"wbtab_{key}", use_container_width=True):
                st.session_state.workbench_tab = key
                st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)

    # Main content: show active job + agent context, then chat
    active_agent = st.session_state.workbench_tab
    prep_job = get_job(st.session_state.prep_job_id) or {}
    active_icon, active_name, _ = AGENT_META[active_agent]
    job_label = f"{prep_job.get('company', '—')} — {prep_job.get('role', '—')}"
    st.markdown(f'<div style="font-size:13px;color:#7a7a8c;font-family:JetBrains Mono,monospace;margin-bottom:12px;">{active_icon} {active_name} · {job_label}</div>', unsafe_allow_html=True)

    render_agent_chat(active_agent, show_job_selector=False, chat_input_key=f"wb_{active_agent}", skip_chat_input=True)

    # Single chat input for the workbench, routed to active tab's agent
    active_name = AGENT_META[active_agent][1]
    wb_input = st.chat_input(f"Ask {active_name}...", key="workbench_chat_input")
    if wb_input and wb_input.strip():
        msgs = _get_prep_messages(active_agent, st.session_state.prep_job_id)
        msgs.append({"role": "user", "content": wb_input.strip()})
        _set_prep_messages(active_agent, st.session_state.prep_job_id, msgs)
        st.rerun()

# ─── PROFILE ───
elif st.session_state.current_page == "profile":
    st.title("My Profile")
    st.caption("Your profile powers all AI agents. The more detail you provide, the better the coaching.")

    p = st.session_state.profile

    col1, col2 = st.columns(2)
    with col1:
        name = st.text_input("Your Name", value=p.get("name", ""), placeholder="e.g. Annie Chen")
    with col2:
        target_role = st.text_input("Target Role", value=p.get("role", ""), placeholder="e.g. ML Engineer Intern, Data Science Intern")

    st.markdown("**Internship Context** — used by all agents to calibrate advice")
    col_uni, col_grad, col_gpa = st.columns(3)
    with col_uni:
        university = st.text_input("University", value=p.get("university", ""), placeholder="e.g. NUS, NTU, SMU")
    with col_grad:
        graduation = st.text_input("Graduation (cycle)", value=p.get("graduation", ""), placeholder="e.g. May 2026 / Summer 2026")
    with col_gpa:
        gpa = st.text_input("GPA", value=p.get("gpa", ""), placeholder="e.g. 4.2/5.0 or 3.8/4.0")

    resume = st.text_area(
        "Resume / Background (paste full text)",
        value=p.get("resume", ""),
        height=220,
        placeholder="Your work experience, education, skills, projects...\n\nThe more detail, the more personalized your coaching will be."
    )

    goals = st.text_area(
        "Goals & Current Situation",
        value=p.get("goals", ""),
        height=100,
        placeholder="e.g. Landing a summer 2026 SWE or DS intern offer. Targeting MNCs and top startups in Singapore/US. Unsure whether to go SWE vs DS track."
    )

    resume_constraints = st.text_input(
        "Resume Space Constraints",
        value=p.get("resume_constraints", ""),
        placeholder="e.g. 2 work experiences, 3 projects, max 1 page",
        help="The Resume Reviewer enforces these limits every session — you never need to repeat them."
    )

    if st.button("💾 Save Profile", use_container_width=False):
        st.session_state.profile = {
            "name": name, "role": target_role,
            "university": university, "graduation": graduation, "gpa": gpa,
            "resume": resume, "goals": goals,
            "resume_constraints": resume_constraints
        }
        db_save("profile", st.session_state.profile)
        st.success("✓ Profile saved! Your agents are now fully personalized.")

    st.divider()
    st.subheader("📅 Google Calendar")
    st.caption("Connect your calendar so the Career Coach can plan your schedule — block study time, add interview prep, and suggest when to work on applications.")
    if is_calendar_connected():
        st.success("✓ Google Calendar connected. The Career Coach can view and create events.")
        if st.button("Disconnect Google Calendar", key="disconnect_cal"):
            disconnect_calendar()
            st.rerun()
    else:
        if st.button("Connect Google Calendar", key="connect_cal"):
            with st.spinner("Opening browser for Google sign-in..."):
                ok, msg = connect_calendar_oauth()
            if ok:
                st.success(msg)
                st.rerun()
            else:
                st.error(msg)

    st.divider()
    st.subheader("Quick Tips")
    st.markdown("""
    **For best results:**
    - **Resume**: Paste your actual resume text (not formatted) — experience, education, projects, skills
    - **Goals**: What role, what timeline, what's your biggest uncertainty right now?
    - **Job Descriptions**: Paste JDs directly into each job in the tracker — this powers all 6 agents
    - **Study Partner & Planner**: Already know how you learn best — output-first, builder-strategist style is baked in
    """)

# ─── PROJECT LIBRARY ───
elif st.session_state.current_page == "project_library":
    st.title("📚 Project Library")
    st.caption("Upload your project reports or paste descriptions. Claude extracts verified facts that power the Resume Reviewer — so it rewrites bullets with YOUR real data, never hallucinated details.")

    projects = db_get_projects()

    # ── Status bar ──
    if projects:
        st.markdown(f"""<div style="background:#101115;border:1px solid rgba(124,106,247,0.25);border-radius:8px;
        padding:8px 14px;font-family:'JetBrains Mono',monospace;font-size:12px;color:#7a7a8c;margin-bottom:16px;">
        📚 {len(projects)} project{'s' if len(projects)!=1 else ''} in library &nbsp;·&nbsp;
        🔧 {sum(len(p['technologies']) for p in projects)} technologies indexed &nbsp;·&nbsp;
        📊 {sum(len(p['metrics']) for p in projects)} metrics captured
        </div>""", unsafe_allow_html=True)

    # ── Add New Project ──
    with st.expander("➕  Add a New Project", expanded=not bool(projects)):
        st.markdown("**Upload a report or paste your project description**")
        st.caption("Supported: PDF, Word (.docx), Markdown (.md), plain text — or just type/paste directly.")

        tab_pdf, tab_docx, tab_file, tab_text = st.tabs(["📄 PDF", "📝 Word (.docx)", "🗒️ Markdown / Text", "✏️ Paste / Type"])

        raw_content = ""
        source_type = "text"

        with tab_pdf:
            pdf_file = st.file_uploader("Upload PDF project report", type=["pdf"], key="proj_pdf")
            if pdf_file:
                with st.spinner("Extracting text from PDF..."):
                    raw_content = read_pdf_bytes(pdf_file.read())
                source_type = "pdf"
                if raw_content and not raw_content.startswith("[PDF extraction"):
                    st.success(f"✓ Extracted ~{len(raw_content.split())} words from PDF")
                    with st.expander("Preview extracted text"):
                        st.text(raw_content[:1500] + ("..." if len(raw_content) > 1500 else ""))
                else:
                    st.warning(raw_content or "Could not extract text from this PDF.")

        with tab_docx:
            docx_file = st.file_uploader("Upload Word document (.docx)", type=["docx"], key="proj_docx")
            if docx_file:
                with st.spinner("Extracting text from Word document..."):
                    raw_content = read_docx_bytes(docx_file.read())
                source_type = "docx"
                if raw_content and not raw_content.startswith("[DOCX extraction"):
                    st.success(f"✓ Extracted ~{len(raw_content.split())} words from Word doc")
                    with st.expander("Preview extracted text"):
                        st.text(raw_content[:1500] + ("..." if len(raw_content) > 1500 else ""))
                else:
                    st.warning(raw_content or "Could not extract text from this Word document.")

        with tab_file:
            md_file = st.file_uploader("Upload .md or .txt file", type=["md", "txt"], key="proj_md")
            if md_file:
                raw_content = md_file.read().decode("utf-8", errors="replace")
                source_type = "markdown" if md_file.name.endswith(".md") else "text"
                st.success(f"✓ Loaded ~{len(raw_content.split())} words")
                with st.expander("Preview"):
                    st.text(raw_content[:1500] + ("..." if len(raw_content) > 1500 else ""))

        with tab_text:
            pasted = st.text_area(
                "Paste or type your project description",
                height=200,
                placeholder="Describe your project: what you built, tech stack, your role, metrics, challenges...\n\nYou can paste a report, README, project summary, or write it yourself.",
                key="proj_paste"
            )
            if pasted.strip():
                raw_content = pasted.strip()
                source_type = "text"

        st.divider()

        title_hint = st.text_input(
            "Project title (optional — Claude will infer if blank)",
            placeholder="e.g. Fraud Detection Pipeline, RAG-based Study Assistant...",
            key="proj_title_hint"
        )

        col_add, col_space = st.columns([2, 3])
        with col_add:
            add_btn = st.button("🔍 Extract & Add to Library", use_container_width=True, key="add_project_btn",
                                disabled=not bool(raw_content))

        if add_btn and raw_content:
            with st.spinner("Claude is reading your project and extracting facts..."):
                extracted = extract_project_with_claude(raw_content, title_hint.strip())
            final_title = extracted.get("title") or title_hint.strip() or "Untitled Project"
            db_save_project(final_title, source_type, raw_content, extracted)
            st.success(f"✅ **{final_title}** added to your library!")
            st.markdown(f"""
<div style="background:#101115;border:1px solid rgba(124,106,247,0.2);border-radius:8px;padding:14px 18px;margin-top:8px;font-size:13px;font-family:'JetBrains Mono',monospace;">
<span style="color:#7c6af7">🔧 Tech:</span> <span style="color:#eeedf0">{', '.join(extracted.get('technologies', [])) or '—'}</span><br>
<span style="color:#7c6af7">📊 Metrics:</span> <span style="color:#eeedf0">{' | '.join(extracted.get('metrics', [])) or 'None found in text'}</span><br>
<span style="color:#7c6af7">🧑 Contributions:</span> <span style="color:#eeedf0">{extracted.get('contributions', '')[:200] or '—'}</span>
</div>""", unsafe_allow_html=True)
            st.rerun()

    st.divider()

    # ── Library Display ──
    if not projects:
        st.markdown("""<div style="text-align:center;padding:48px 20px;color:#7a7a8c;">
        <div style="font-size:40px;margin-bottom:12px;">📂</div>
        <div style="font-size:15px;font-weight:600;margin-bottom:6px;">No projects yet</div>
        <div style="font-size:13px;">Add a project above — the Resume Reviewer will use your real facts to write stronger, accurate bullets.</div>
        </div>""", unsafe_allow_html=True)
    else:
        st.subheader(f"Your Library  ·  {len(projects)} project{'s' if len(projects)!=1 else ''}")
        st.caption("These facts are automatically injected into the Resume Reviewer every time it runs.")

        for proj in projects:
            created_str = datetime.fromtimestamp(proj["created_at"]).strftime("%b %d, %Y")
            source_icon = {"pdf": "📄", "docx": "📝", "markdown": "🗒️", "text": "✏️"}.get(proj["source_type"], "📁")

            with st.expander(f"{source_icon}  **{proj['title']}**  —  added {created_str}"):
                col_l, col_r = st.columns([4, 1])

                with col_l:
                    if proj["technologies"]:
                        st.markdown("**🔧 Technologies**")
                        tech_badges = "  ".join(
                            [f'<span style="background:rgba(124,106,247,0.1);border:1px solid rgba(124,106,247,0.25);'
                             f'border-radius:4px;padding:2px 8px;font-size:12px;font-family:\'JetBrains Mono\',monospace;'
                             f'color:#7c6af7">{t}</span>'
                             for t in proj["technologies"]]
                        )
                        st.markdown(tech_badges, unsafe_allow_html=True)
                        st.markdown("")

                    if proj["metrics"]:
                        st.markdown("**📊 Metrics & Outcomes**")
                        for m in proj["metrics"]:
                            st.markdown(f"- {m}")

                    if proj["contributions"]:
                        st.markdown("**🧑 Your Contributions**")
                        st.markdown(proj["contributions"])

                    if proj["challenges"]:
                        st.markdown("**⚙️ Challenges Solved**")
                        st.markdown(proj["challenges"])

                    if proj["summary"]:
                        st.markdown("**💡 Summary**")
                        st.markdown(f"_{proj['summary']}_")

                with col_r:
                    st.markdown("<br>", unsafe_allow_html=True)
                    if st.button("🗑 Delete", key=f"del_proj_{proj['id']}", use_container_width=True):
                        db_delete_project(proj["id"])
                        st.success(f"Deleted '{proj['title']}'")
                        st.rerun()

        st.divider()
        st.markdown(f"""<div style="background:#101115;border:1px solid rgba(255,255,255,0.06);border-radius:8px;
        padding:12px 16px;font-size:12px;color:#7a7a8c;font-family:'JetBrains Mono',monospace;">
        💡 <strong style="color:#eeedf0">How this powers Resume Reviewer:</strong>
        The Resume Reviewer reads your full project library before every session. It auto-matches your projects to the
        target JD and uses your real technologies, metrics, and contributions when rewriting bullets —
        no hallucinations, no underestimating what you actually built.
        </div>""", unsafe_allow_html=True)
